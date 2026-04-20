from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor, Cm

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.header_section.header_section import \
    header_section
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.experiences.display_environnement_entreprise import \
    display_environnement_entreprise
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.experiences.logiciels_outils.display_logiciels_outils import \
    display_logiciels_outils
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.experiences.poste.display_poste import \
    display_poste
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.experiences.mission.display_mission import display_mission
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.experiences.tasks.display_tasks import \
    display_tasks


async def build_section_experiences(doc, data):

    header_section(doc, "EXPERIENCES PROFESSIONNELLES")

    for exp in data.get('Expériences_Professionnelles_Antéchronologiques', []):
        p = doc.add_paragraph()
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True

        tab_xml = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="center" w:pos="5125"/><w:tab w:val="right" w:pos="9639"/></w:tabs>')
        p._element.get_or_add_pPr().append(tab_xml)

        border_xml = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:color="002060"/></w:pBdr>')
        p._element.get_or_add_pPr().append(border_xml)


        run_entreprise = p.add_run(exp.get('Nom_Entreprise', '').upper())
        run_entreprise.bold = True
        run_entreprise.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        p.add_run('\t')

        run_duree = p.add_run(f"({exp.get('Durée_Mission_Ou_Mission_Toujours_En_Cours', '')})")
        run_duree.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        p.add_run('\t')

        run_dates = p.add_run(str(exp.get('Dates_Période', '')))
        run_dates.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        if exp.get("Poste_Occupé"):
            display_poste(doc, exp)

        # if exp.get('Phrase_Courte_Expliquant_Les_Secteurs_Activité_Entreprise'):
        await display_environnement_entreprise(doc, exp)

        if exp.get('Résumé_mission_En_Phrase_Nominal'):
            display_mission(doc, exp)

        if exp.get('Liste_Tâches_En_Phrase_Nominal'):
            display_tasks(doc, exp)

        if exp.get("Logiciels_et_outils_utilisés_Sans_Indiquer_Le_Niveau"):
            display_logiciels_outils(doc, exp)
