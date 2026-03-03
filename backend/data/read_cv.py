# import os

# import fitz
# from docx import Document


# def read_cv(file_path):
#     doc = fitz.open(file_path)
#     ext = os.path.splitext(file_path)[1].lower()
#     text = ""

#     if ext == ".pdf":
#         with fitz.open(file_path) as doc:
#             for page in doc:
#                 blocks = page.get_text("blocks", sort=True)
#                 for block in blocks:
#                     text += block[4] + "\n"

#     elif ext == ".docx":
#         doc = Document(file_path)
#         for para in doc.paragraphs:
#             if para.text:
#                 text += para.text + "\n"

#         for table in doc.tables:
#             for row in table.rows:
#                 text += " | ".join(cell.text for cell in row.cells) + "\n"

#     else:
#         return "Format non supporté (utilisez PDF ou DOCX)"

#     # text = text[:25000]
#     print(text)
#     return text


# import io
import os

import cv2
import fitz  # PyMuPDF
import numpy as np
import pytesseract
from docx import Document

# from PIL import Image


def read_cv(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        with fitz.open(file_path) as doc:
            for page in doc:
                page_dict = page.get_text("dict", sort=True)
                if isinstance(page_dict, dict) and "blocks" in page_dict:
                    for block in sorted(
                        page_dict.get("blocks", []),
                        key=lambda b: (int(b["bbox"][1] / 5), b["bbox"][0]),
                    ):
                        if isinstance(block, dict) and "lines" in block:
                            # if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = f"{text}{span['text']} "
                            text += "\n"

            if len(text.strip()) < 200:
                print("📸 PDF image détecté, lancement de l'OCR...")
                text = ""
                for page in doc:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                        pix.h, pix.w, pix.n
                    )
                    img = cv2.resize(
                        img, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC
                    )
                    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                    thresh = cv2.adaptiveThreshold(
                        gray,
                        255,
                        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                        cv2.THRESH_BINARY,
                        11,
                        10,
                    )
                    text += (
                        pytesseract.image_to_string(
                            thresh,
                            lang="fra+eng",
                            config="--psm 3",
                        )
                        + "\n"
                    )

    elif ext == ".docx":
        doc = Document(file_path)
        text += "\n".join([p.text for p in doc.paragraphs if p.text])
        for table in doc.tables:
            for row in table.rows:
                text += " | ".join(c.text.strip() for c in row.cells) + "\n"

    print(f"✅ Extraction terminée : {len(text)} caractères trouvés.")
    print(text)
    return text
