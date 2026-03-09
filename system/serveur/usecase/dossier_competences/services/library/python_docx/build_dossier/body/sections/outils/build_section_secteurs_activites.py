from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

def build_section_outils(doc, data):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("OUTILS")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True

    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002060"/>')
    p1._element.get_or_add_pPr().append(shd)

    for categorie in data.get('Logiciels_par_titre', []):
        print(categorie)
        titre = categorie.get('titre', '')
        if titre:
            p_titre = doc.add_paragraph(titre)
            p_titre.runs[0].font.bold = True
        for outil in categorie.get('logiciels_outils', []):
            print(outil)
            p_item = doc.add_paragraph(outil, style='List Bullet')
            p_item.paragraph_format.space_after = Pt(2)
            p_item.paragraph_format.keep_together = True