from docx.shared import Pt, RGBColor


def display_mission(doc, exp):
    p_mission = doc.add_paragraph()

    run_t = p_mission.add_run("MISSION")
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_t.bold = True
    run_t.underline = True

    run_sep = p_mission.add_run(" : ")
    run_sep.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_sep.bold = True

    run_desc = p_mission.add_run(str(exp.get('Résumé_mission', '')))
    run_desc.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_mission.paragraph_format.space_before = Pt(10)
    p_mission.paragraph_format.space_after = Pt(10)