from docx.shared import RGBColor, Cm


def left_cells(cells, diplome):
    p = cells[0].paragraphs[0]
    p.paragraph_format.keep_with_next = True

    run_dip = p.add_run(diplome.get('Diplôme', ''))
    run_dip.bold = True
    run_dip.underline = True
    run_dip.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    if diplome.get("École"):
        p_info = cells[0].add_paragraph()
        p_info.paragraph_format.left_indent = Cm(1.25)
        texte = diplome.get("École")
        if diplome.get("Lieu"):
            texte += f" - {diplome.get('Lieu')}"

        run_info = p_info.add_run(texte)
        run_info.font.color.rgb = RGBColor(0x00, 0x20, 0x60)