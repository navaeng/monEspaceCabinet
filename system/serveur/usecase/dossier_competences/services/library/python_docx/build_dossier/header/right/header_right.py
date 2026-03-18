from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from usecase.dossier_competences.services.library.python_docx.build_dossier.header.right.generate.generate_initial import \
    generate_initial
from usecase.dossier_competences.services.library.python_docx.build_dossier.header.right.generate.generate_exp import \
    generate_exp

POLICE = "Calibri (corps)"

def header_right(doc, data, p):

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    initiales = generate_initial(data)
    run1 = p.add_run(initiales)
    run1.font.name = POLICE
    run1.font.bold = False
    run1.font.size = Pt(17)
    run1.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    run2 = p.add_run(data.get('Poste_actuel', ''))
    run2.font.name = POLICE
    run2.font.size = Pt(20)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    run3 = p.add_run(generate_exp(data))
    run3.font.name = POLICE
    run3.font.bold = False
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x00, 0x20, 0x60)