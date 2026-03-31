import json

from docx.shared import Pt, RGBColor

from services.api_externes.openrouter.faire_recherche_openrouter import faire_recherche_openrouter


# def display_environnement_entreprise(doc, exp):
#     p_mission = doc.add_paragraph()
#     p_mission.paragraph_format.keep_together = True
#     p_mission.paragraph_format.keep_with_next = True
#
#     run_t = p_mission.add_run("Environnement")
#     run_t.font.size = Pt(11)
#     run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
#     run_t.bold = True
#
#     run_sep = p_mission.add_run(" : ")
#     run_sep.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
#     run_sep.bold = True
#
#     run_desc = p_mission.add_run(str(exp.get('Phrase_Courte_Expliquant_Les_Secteurs_Activité_Entreprise', '')))
#     run_desc.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
#
#     p_mission.paragraph_format.space_before = Pt(10)
#     p_mission.paragraph_format.space_after = Pt(10)


async def display_environnement_entreprise(doc, exp):

    nom_boite = exp.get('Nom_Entreprise', '')
    print(f"DEBUG: Recherche pour {nom_boite}...")

    secteur_reel = await faire_recherche_openrouter(nom_boite)
    print(f"DEBUG: Résultat IA pour {nom_boite} -> '{secteur_reel}'")

    try:
        secteur_reel = json.loads(secteur_reel.replace("```json", "").replace("```", "").strip()).get("Secteur_entreprise",
                                                                                                      secteur_reel)
    except (json.JSONDecodeError, AttributeError):
        pass

    p_mission = doc.add_paragraph()
    p_mission.paragraph_format.space_before = Pt(10)
    p_mission.paragraph_format.space_after = Pt(10)

    run_t = p_mission.add_run("Environnement")
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_t.bold = True

    run_sep = p_mission.add_run(" : ")
    run_sep.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_sep.bold = True

    run_desc = p_mission.add_run(str(secteur_reel))
    run_desc.font.color.rgb = RGBColor(0x00, 0x20, 0x60)