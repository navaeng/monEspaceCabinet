import os

import fitz
from docx import Document


def read_cv(file_path):
    doc = fitz.open(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        with fitz.open(file_path) as doc:
            for page in doc:
                blocks = page.get_text("blocks", sort=True)
                for block in blocks:
                    text += block[4] + "\n"

    elif ext == ".docx":
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                text += " | ".join(cell.text for cell in row.cells) + "\n"

    else:
        return "Format non supporté (utilisez PDF ou DOCX)"

    print(text)
    return text
