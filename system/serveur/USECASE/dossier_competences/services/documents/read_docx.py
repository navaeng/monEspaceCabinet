from docx import Document


def read_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text])
    for table in doc.tables:
        for row in table.rows:
            text += " | ".join(c.text.strip() for c in row.cells) + "\n"
    return text