from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.header_section.header_section import \
    header_section


def build_section_outils(doc, data):
    outils_data = data.get('Logiciels_Et_Outils_Sans_Indiquer_Le_Niveau', [])
    if outils_data and outils_data[0].get('Liste_Logiciels') and len(outils_data[0]['Liste_Logiciels']) > 0:

        header_section(doc, "LOGICIELS ET OUTILS")

        for categorie in data.get('Logiciels_Et_Outils_Sans_Indiquer_Le_Niveau', []):
            print(categorie)
            titre = categorie.get('Catégorie', '')
            logiciels_outils = list(dict.fromkeys(categorie.get('Liste_Logiciels', [])))

           if titre and logiciels_outils:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.keep_together = True

                run_t = p.add_run(f"{titre.upper()} : ")
                run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                run_t.bold = True
                run_t.underline = True

                run_l = p.add_run(", ".join(logiciels_outils))
                run_l.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

                p.paragraph_format.space_after = Pt(4)
