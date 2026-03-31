from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor, Cm  # Ajout de Cm ici

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.header_section.header_section import \
    header_section


def build_section_langues(doc, data):
    langues = data.get('Langues', [])

    if not langues or not langues[0].get('Langue'):
        return

    header_section(doc, "LANGUES")

    for langue in data.get('Langues', []):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.keep_together = True
        
        # On définit un taquet de tabulation (ex: à 4cm du bord gauche)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4))

        run_l = p.add_run(langue.get('Langue'))
        run_l.bold = True
        run_l.font.color.rgb = RGBColor(0, 20, 60)

        run_n = p.add_run(f"\t{langue.get('Niveau')}")
        run_n.font.color.rgb = RGBColor(0, 20, 60)
        
        p.paragraph_format.space_after = Pt(2)
