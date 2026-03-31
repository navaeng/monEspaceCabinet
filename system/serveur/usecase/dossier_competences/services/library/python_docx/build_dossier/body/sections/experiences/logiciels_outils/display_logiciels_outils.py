from docx.shared import RGBColor, Pt

def display_logiciels_outils(doc, exp):
    p_outils = doc.add_paragraph()
    p_outils.paragraph_format.keep_together = True
    p_outils.paragraph_format.keep_with_next = False


    liste_outils = exp.get('Logiciels_et_outils_utilisés_Sans_Indiquer_Le_Niveau', [])
    texte_outils = ", ".join(liste_outils) if isinstance(liste_outils, list) else str(liste_outils)

    run_t = p_outils.add_run("Logiciels et outils")
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_t.bold = True

    run_sep = p_outils.add_run(" : ")
    run_sep.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_sep.bold = True

    run_liste = p_outils.add_run(texte_outils)
    run_liste.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_outils.paragraph_format.space_before = Pt(6)
    p_outils.paragraph_format.space_after = Pt(20)

    p_outils.paragraph_format.keep_with_next = False
    p_outils.paragraph_format.keep_together = True
