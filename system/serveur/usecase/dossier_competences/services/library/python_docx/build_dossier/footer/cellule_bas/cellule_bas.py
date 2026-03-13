from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt


def cellule_bas(tab):
    p2 = tab.cell(1, 0).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p2.add_run("NAVA ENGINEERING • 9-11 AVENUE MICHELET, 93400 SAINT-OUEN-SUR-SEINE\n"
                    "Simon ZANA • 06 13 53 23 81 • ")
    r1.font.size, r1.font.color.rgb = Pt(9), RGBColor(0, 32, 96)

    r2 = p2.add_run("simon.zana@nava-eng.com")
    r2.font.size, r2.font.color.rgb, r2.font.underline = Pt(8), RGBColor(0, 32, 96), True


