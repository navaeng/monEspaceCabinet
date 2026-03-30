from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Cm, Pt

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.header_section.header_section import \
    header_section
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.formations.cells.left_cells import \
    left_cells
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.formations.cells.right_cells import \
    right_cells


def build_section_formation(doc, data):

    formations = data.get('Diplômes_Et_Formations_Antéchronologiques') or []

    if not formations or not formations[0].get('Diplôme'):
        return

    header_section(doc, "FORMATIONS/DIPLOMES")

    for diplome in data.get('Diplômes_Et_Formations_Antéchronologiques', []):
        p = doc.add_paragraph(diplome.get('Diplôme', ''), style='List Bullet')
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_after = Pt(3.5)

        tab_xml = parse_xml(f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:pos="9639"/></w:tabs>')
        p._element.get_or_add_pPr().append(tab_xml)

        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        p.add_run('\t')

        run_date = p.add_run(str(diplome.get('Année', '')))
        run_date.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        if diplome.get('École'):

            p_ecole = doc.add_paragraph()
            p_ecole.paragraph_format.left_indent = Cm(1.25)
            p_ecole.paragraph_format.space_after = Pt(0.5)
            p_ecole.paragraph_format.keep_with_next = True
            run_ecole = p_ecole.add_run(diplome.get('École'))
            run_ecole.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

            if diplome.get('Lieu'):
                p_lieu = doc.add_paragraph()
                p_lieu.paragraph_format.left_indent = Cm(1.25)
                run_lieu = p_lieu.add_run(diplome.get('Lieu'))
                run_lieu.font.color.rgb = RGBColor(0x00, 0x20, 0x60)