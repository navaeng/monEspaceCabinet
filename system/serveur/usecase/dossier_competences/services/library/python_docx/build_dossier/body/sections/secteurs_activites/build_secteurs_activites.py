from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

def build_section_secteurs_activites(doc, data):
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("SECTEURS ACTIVITES")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True

    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002060"/>')
    p2._element.get_or_add_pPr().append(shd)

    for secteur in data.get('Secteurs_activités', []):
        p_item = doc.add_paragraph(secteur, style='List Bullet')
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.left_indent = Pt(20)
        p_item.paragraph_format.keep_together = True
        p_item.paragraph_format.keep_with_next = True
        print(secteur)