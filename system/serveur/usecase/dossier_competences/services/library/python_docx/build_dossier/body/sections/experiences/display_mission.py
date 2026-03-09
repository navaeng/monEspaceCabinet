from docx.shared import Pt

def display_mission(doc, exp):
    p_mission = doc.add_paragraph()

    p_mission.add_run("Mission : ").bold = True
    p_mission.add_run(str(exp.get('Mission', '')))

    p_mission.paragraph_format.space_before = Pt(10)
    p_mission.paragraph_format.space_after = Pt(10)
