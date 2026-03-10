from docx.shared import RGBColor, Pt

def display_logiciels_outils(doc, exp):
    p_outils = doc.add_paragraph()

    liste_outils = exp.get('Environnement_Technique', [])
    texte_outils = ", ".join(liste_outils) if isinstance(liste_outils, list) else str(liste_outils)

    run_label = p_outils.add_run("LOGICIELS ET OUTILS : ")
    run_label.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run_label.bold = True

    run_liste = p_outils.add_run(texte_outils)
    run_liste.font.color.rgb = RGBColor(0, 0, 0)

    p_outils.paragraph_format.space_before = Pt(6)
    p_outils.paragraph_format.space_after = Pt(20)

    p_outils.paragraph_format.keep_with_next = False
    p_outils.paragraph_format.keep_together = True
